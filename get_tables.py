#!/usr/bin/env python3
# -*- encoding: utf-8 -*-

from docx import Document

wordDoc = Document("./sample.docx")

# 表の数を出力
print(f"num of tables = {len(wordDoc.tables)}")

for table in wordDoc.tables:
    # 表の行数を出力
    print(f"num of rows = {len(table.rows)}")

    # 表のヘッダ行を出力
    print(f"header of table is {str(table.rows[0])}")
    for cell in table.rows[0].cells:
        print(cell.text)

    # 表中の全セルを出力
    for row in table.rows:
        print(f"num of cell = {len(row.cells)}")
        for cell in row.cells:
            print(cell.text)

